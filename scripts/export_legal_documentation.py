from __future__ import annotations

from pathlib import Path
from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = PROJECT_ROOT / "PROJECT_DOCUMENTATION.md"
OUTPUT_DOCX = PROJECT_ROOT / "LegalAI_Documentation.docx"


def main():
    document = Document()

    if not SOURCE_MD.exists():
        print("❌ Markdown file not found.")
        return

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()

    for line in lines:
        if line.startswith("#"):
            document.add_heading(line.replace("#", "").strip(), level=1)
        else:
            document.add_paragraph(line)

    document.save(OUTPUT_DOCX)
    print(f" DOCX exported → {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()